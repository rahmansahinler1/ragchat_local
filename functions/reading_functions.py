import PyPDF2
from docx import Document
import spacy
from pathlib import Path
from datetime import datetime
import os 
import fitz
import re
import pymupdf4llm
from langchain_text_splitters import MarkdownHeaderTextSplitter

class ReadingFunctions:
    def __init__(self):
        self.nlp = spacy.load(
            "en_core_web_sm",
            disable=[ "tagger", "attribute_ruler", "lemmatizer", "ner","textcat","custom "]
        )
        self.headers_to_split_on = [
            ("#", "Header 1"),
            ("##", "Header 2"),
            ("###", "Header 3"),
            ("####", "Header 4")
        ]
        self.markdown_splitter = MarkdownHeaderTextSplitter(self.headers_to_split_on,strip_headers=False,return_each_line=True)
    def read_file(self, file_path: str):
        file_data = {
            "page_sentence_amount": [],
            "sentences": [],
            "date": [],
            "is_header": [],
            "page_num": [],
            "file_header" : [],
            "is_table" : [],
        }
        # Open file
        path = Path(file_path)
        file_extension = path.suffix.lower()
        try:
            if file_extension == '.pdf':
                markdown_pages = pymupdf4llm.to_markdown(path, page_chunks=True)
                with fitz.open(path) as file:
                    try:
                        pdf_date = f"{file.metadata["creationDate"][4:6]}-{file.metadata["creationDate"][6:8]}-{file.metadata["creationDate"][9:11]}"
                        file_data["date"].append(pdf_date)
                        previous_sentence_count = 0
                    except TypeError as e:
                        raise TypeError(f"PDF creation date could not extracted!: {e}")
                    try: 
                        for i,page in enumerate(markdown_pages):
                            splits = self.markdown_splitter.split_text(page['text'])
                            for split in splits:
                                if not len(split.page_content) > 5 or re.match(r'^[^\w]*$',split.page_content) or re.match(r'E\/ECE\/(?:\d+|TRANS)\/Rev\.\d+\/Add\.\d+\/Rev\.\d+',split.page_content):
                                    continue
                                elif split.metadata and split.page_content[0] == '#' : #header
                                    file_data['sentences'].append(split.page_content)
                                    file_data['is_header'].append(1)
                                    file_data['is_table'].append(0)
                                    file_data['page_num'].append(i+1)
                                elif split.page_content[0] == "*" and split.page_content[-1] == '*' and (re.match(r"(\*{2,})(\d+(?:\.\d+)*)\s*(\*{2,})?(.*)$",split.page_content) or re.match(r"(\*{1,3})?([A-Z][a-zA-Z\s\-]+)(\*{1,3})?$",split.page_content)): # Sub-Header and Header variant detection
                                    file_data["sentences"].append(split.page_content)
                                    file_data["is_header"].append(1)
                                    file_data["is_table"].append(0)
                                    file_data["page_num"].append(i+1)
                                elif split.page_content[0] == '|' and split.page_content[-1] == '|': #table
                                    file_data['sentences'].append(split.page_content)
                                    file_data['is_header'].append(0)
                                    file_data['is_table'].append(1)
                                    file_data['page_num'].append(i+1)
                                else:
                                    file_data['sentences'].append(split.page_content)
                                    file_data['is_header'].append(0)
                                    file_data['is_table'].append(0)
                                    file_data['page_num'].append(i+1)
                            current_sentence_count = len(file_data["sentences"])
                            sentences_in_this_page = current_sentence_count - previous_sentence_count
                            file_data["page_sentence_amount"].append(sentences_in_this_page)
                            previous_sentence_count = current_sentence_count
                    except TypeError as e:
                        raise TypeError(f"PDF info could not extracted!: {e}")
            elif file_extension == '.docx':
                doc = Document(path)
                try:
                    previous_sentence_count = 0
                    doc_date = f"{doc.core_properties.created.year%2000}-{doc.core_properties.created.month}-{doc.core_properties.created.day}"
                    file_data["date"].append(doc_date)
                except TypeError as e:
                    raise TypeError(f"Doc creation date could not extracted!: {e}")
                finally:
                    file_data["date"].append('')
                    try:
                        for para in doc.paragraphs:
                            if para.style.name.startswith('Heading') or para.style.name.startswith('Title'):
                                file_data["sentences"].append(para.text)
                                file_data["is_header"].append(1)
                                file_data["is_table"].append(0)
                            elif len(para.text) > 15:
                                file_data["sentences"].append(para.text)
                                file_data["is_header"].append(0)
                                file_data["is_table"].append(0)
                        current_sentence_count = len(file_data["sentences"])
                        sentences_in_this_page = current_sentence_count - previous_sentence_count
                        file_data["page_sentence_amount"].append(sentences_in_this_page)
                        previous_sentence_count = current_sentence_count
                    except TypeError as e:
                            raise TypeError(f"PDF text could not extracted!: {e}")
            elif file_extension in ['.txt', '.rtf']:
                txt_date = os.path.getctime(path)
                file_data["date"].append(txt_date)
                text = path.read_text(encoding='utf-8')
                self._process_text(text, file_data)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        
        except PyPDF2.errors.PdfReadError:
            print(f"Error reading PDF file: {path}. The file might be corrupted or incompatible.")
        except Exception as e:
            print(f"Error reading file: {path}. Error: {str(e)}")
        return file_data
    
    def _process_text(self, text, file_data):
       docs = self.nlp(text)
       sentences = [sent.text.replace('\n', ' ').strip() for sent in docs.sents]
       valid_sentences = [sentence for sentence in sentences if len(sentence) > 15]
       file_data["page_sentence_amount"].append(len(valid_sentences))
       file_data["sentences"].extend(valid_sentences)

        # Extract files first page header  
    def _extract_file_header(self, file_path, file_data):
        path = Path(file_path)
        file_extension = path.suffix.lower()
        full_text = ''
        try:
            if file_extension == '.pdf':
                doc = fitz.open(path)
                page = doc.load_page(0)
                blocks = page.get_text("dict")["blocks"]
                text_blocks = [block for block in blocks if block["type"] == 0]
                for block in text_blocks:
                    if "lines" in block and len(block["lines"]) >= 1 and len(block["lines"]) < 4:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"]
                                if span["size"] > 3 and (span["font"].find("Medi") >0 or span["font"].find("Bold") >0 or span["font"].find("B") >0) and len(text) > 3:
                                    full_text += text + ' '
                file_data["file_header"].append(full_text)
            elif file_extension == '.docx':
                doc = Document(path)
                for para in doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        full_text += text + ' '
                file_data["file_header"].append(full_text)
        except Exception as e:
            print(f"Error reading file: {path}. Error: {str(e)}")

